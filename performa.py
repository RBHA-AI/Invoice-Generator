import os
import re
from decimal import Decimal, ROUND_HALF_UP
import pandas as pd
from docx import Document

EXCEL_PATH = "calculation.xlsx"
TEMPLATE_DOCX = "proforma Sonepat 20_.docx"
OUTPUT_DIR = "output_proformas"

HSN_CODE = "998222"

# anchors present in template
OLD_BUYER_NAME = "Municipal Corporation Sonepat"
OLD_LOCATION_WORD = "Sonepat"

def D(x) -> Decimal:
    return Decimal(str(x))

def money_2(x: Decimal) -> str:
    return str(x.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))

def sanitize_filename(s: str) -> str:
    s = str(s).strip()
    s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
    s = re.sub(r"\s+", " ", s)
    return s

def make_buyer_name(location: str) -> str:
    location = str(location).strip()
    if location.lower().startswith("municipal"):
        return location
    return f"Municipal Corporation {location}"

# ----------------------------
# Amount to words (Indian system)
# ----------------------------
ONES = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
        "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
        "Seventeen", "Eighteen", "Nineteen"]
TENS = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]

def two_digit_words(n: int) -> str:
    if n < 20:
        return ONES[n]
    t, o = divmod(n, 10)
    return (TENS[t] + (" " + ONES[o] if o else "")).strip()

def three_digit_words(n: int) -> str:
    h, r = divmod(n, 100)
    out = ""
    if h:
        out += ONES[h] + " Hundred"
        if r:
            out += " "
    if r:
        out += two_digit_words(r)
    return out.strip()

def int_to_indian_words(n: int) -> str:
    if n == 0:
        return "Zero"
    parts = []
    crore, n = divmod(n, 10_000_000)
    lakh, n = divmod(n, 100_000)
    thousand, n = divmod(n, 1000)
    rem = n
    if crore:
        parts.append(int_to_indian_words(crore) + " Crore")
    if lakh:
        parts.append(two_digit_words(lakh) + " Lakh")
    if thousand:
        parts.append(two_digit_words(thousand) + " Thousand")
    if rem:
        parts.append(three_digit_words(rem))
    return " ".join([p for p in parts if p]).strip()

def money_to_words_inr(amount: Decimal) -> str:
    amt = amount.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rupees = int(amt)
    paise = int((amt - Decimal(rupees)) * 100)
    rw = int_to_indian_words(rupees)
    if paise:
        pw = two_digit_words(paise)
        return f"Indian Rupees {rw} and {pw} Paise Only"
    return f"Indian Rupees {rw} Only"

# ----------------------------
# DOCX helpers
# ----------------------------
def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def set_paragraph_text(p, text: str):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ""

def replace_text_everywhere(doc: Document, old: str, new: str):
    if not old or old == new:
        return
    for p in iter_all_paragraphs(doc):
        if not p.runs:
            continue
        full = "".join(r.text for r in p.runs)
        updated = full.replace(old, new)
        if updated != full:
            set_paragraph_text(p, updated)

def replace_regex_everywhere(doc: Document, pattern: str, repl: str):
    rx = re.compile(pattern)
    for p in iter_all_paragraphs(doc):
        if not p.runs:
            continue
        full = "".join(r.text for r in p.runs)
        updated = rx.sub(repl, full)
        if updated != full:
            set_paragraph_text(p, updated)

def fix_buyer_location_in_buyer_block(doc: Document, new_location: str):
    """
    In your file, location is on its own line under Buyer.
    We only replace it inside the Buyer section, not globally.
    """
    in_buyer_block = False
    for p in doc.paragraphs:
        txt = p.text or ""
        if "Buyer" in txt:
            in_buyer_block = True
            continue
        if in_buyer_block:
            # stop when invoice line items start showing up
            if "Sl" in txt and "Particulars" in txt:
                in_buyer_block = False
                continue
            # replace if this paragraph contains the old location word (even with tabs)
            if OLD_LOCATION_WORD in txt:
                updated = txt.replace(OLD_LOCATION_WORD, new_location)
                if updated != txt:
                    set_paragraph_text(p, updated)

    # also check buyer cell if it is in a table in some versions
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "Buyer" in p.text:
                        # likely the buyer block is around here; replace within this cell
                        for pp in cell.paragraphs:
                            if OLD_LOCATION_WORD in pp.text:
                                set_paragraph_text(pp, pp.text.replace(OLD_LOCATION_WORD, new_location))

def set_hsn_in_tables(doc: Document, hsn: str):
    """
    Put HSN into the actual HSN/SAC column cells (both main items table and tax table).
    Avoid any global \d{6} replacement.
    """
    for t in doc.tables:
        if not t.rows:
            continue
        header = [c.text.strip().replace("\n", " ") for c in t.rows[0].cells]
        if "HSN/SAC" not in header:
            continue
        hsn_idx = header.index("HSN/SAC")

        # write HSN into any data rows where HSN cell is empty or wrong
        for r in t.rows[1:]:
            cells = r.cells
            if hsn_idx >= len(cells):
                continue
            current = cells[hsn_idx].text.strip()
            # heuristic: if row looks like an amount row or taxable row and HSN missing/wrong, set it
            row_text = " ".join(c.text.strip() for c in cells)
            if ("Professional Fees" in row_text) or ("Taxable" in row_text) or ("Integrated Tax" in row_text) or (current == "") or (re.fullmatch(r"\d{6}", current) is not None):
                cells[hsn_idx].text = hsn

def fix_amount_chargeable_words(doc: Document, words: str):
    """
    Keep ONLY the bold/main words line:
    - Set the paragraph containing the label to include the words.
    - Clear any following 'Indian Rupees ...' paragraphs until the tax table starts.
    """
    found = False
    for p in iter_all_paragraphs(doc):
        txt = p.text.strip()
        if "Amount Chargeable (in words)" in txt:
            # Keep label + E&O.E as in template, then ONE words line
            # (newline prevents overlap)
            clean = "Amount Chargeable (in words)\tE. & O.E\n" + words
            set_paragraph_text(p, clean)
            found = True
            continue
        if found:
            # stop once tax table header begins
            if txt.replace(" ", "") == "HSN/SAC" or txt.startswith("HSN/SAC"):
                break
            # clear extra duplicate words lines
            if txt.startswith("Indian Rupees"):
                set_paragraph_text(p, "")

def fix_tax_amount_words(doc: Document, tax_words: str):
    """
    Make Tax Amount in words exactly one line, and clear the stray 'Rupees Only' line.
    """
    after_tax = False
    for p in iter_all_paragraphs(doc):
        txt = p.text.strip()
        if "Tax Amount (in words)" in txt:
            set_paragraph_text(p, "Tax Amount (in words) : " + tax_words)
            after_tax = True
            continue
        if after_tax:
            # clear the stray "Rupees Only" line
            if txt == "Rupees Only" or txt.endswith("Rupees Only"):
                set_paragraph_text(p, "")
                after_tax = False

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    df = pd.read_excel(EXCEL_PATH)
    required = ["MCs Name", "taxable", "gst", "total"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError(f"Missing columns in Excel: {missing}\nFound: {list(df.columns)}")

    generated = 0

    for _, row in df.iterrows():
        loc = row["MCs Name"]
        if pd.isna(loc):
            continue
        loc = str(loc).strip()
        if not loc or loc.lower() == "total":
            continue

        location_word = loc
        buyer_name = make_buyer_name(location_word)

        taxable = D(row["taxable"])
        gst = D(row["gst"])
        total = D(row["total"])

        taxable_s = money_2(taxable)
        gst_s = money_2(gst)
        total_s = money_2(total)
        total_int = str(int(total.quantize(Decimal("1"), rounding=ROUND_HALF_UP)))

        amount_words = money_to_words_inr(total)
        tax_words = money_to_words_inr(gst)

        doc = Document(TEMPLATE_DOCX)

        # Buyer name + buyer location (single word)
        replace_text_everywhere(doc, OLD_BUYER_NAME, buyer_name)
        fix_buyer_location_in_buyer_block(doc, location_word)

        # Set HSN only in HSN column cells
        set_hsn_in_tables(doc, HSN_CODE)

        # Replace amounts using your template anchors (safe)
        replace_regex_everywhere(doc, r"\b672591\.52(?:\.00)?\b", taxable_s)
        replace_regex_everywhere(doc, r"\b121066\.47\b", gst_s)
        replace_regex_everywhere(doc, r"\b793658\.00\b", total_s)
        replace_regex_everywhere(doc, r"\b793658\b", total_int)

        # Fix words blocks (no duplicates, no overlap)
        fix_amount_chargeable_words(doc, amount_words)
        fix_tax_amount_words(doc, tax_words)

        out_file = os.path.join(OUTPUT_DIR, f"Proforma_{sanitize_filename(location_word)}.docx")
        doc.save(out_file)
        generated += 1

    print(f"✅ Generated {generated} DOCX files in: {os.path.abspath(OUTPUT_DIR)}")

if __name__ == "__main__":
    main()
